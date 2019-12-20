# Mohamed Hussein ITEC 1150-60.
# This is my final project code that will generate a recipe book of taco
# recipes including a random picture of taco and three randomly generated recipes


import docx  # import this library to create and update a word document

import requests # to get the url
from PIL import Image, ImageDraw, ImageFont # importing image and font
Image = Image.open('photo-1552332386-f8dd00dc2f85.jpeg') # opening image that i will use for my project
Image.thumbnail((800, 800)) # resizing the image not more or less than 800
Image.save('taco_thumbnails.jpeg') # saving the image
img_draw = ImageDraw.Draw(Image) # drawing a text to the image

font = ImageFont.truetype('DejaVuSans.ttf', 20) # using font color for the text that im using for the image
# so i downloaded Dejavu from the website

img_draw.text([200, 80], 'Random Taco Cookbook', fill='purple', font=font)
# for the text in the picture, im using to draw a text in the image
# the ' Random text cookbook' text [200, 80] and fill purple in the text

Image.save('modified_taco.jpg')
# saving the the taco image and naming it while it shows up in the image

# next step im creating a list to store dictionaries in

taco_book = [] # this list i will be using to create three dictionaries for my taco

for i in range(3): # using a loop to get the result url, for my three tacos recipe
    taco_url = requests.get('https://taco-1150.herokuapp.com/random/?full_taco=true').json()
    taco_book.append(taco_url) # appending the taco_book and the url
    print(taco_book) # printing

# creating a word document
taco_word = docx.Document()
# staring with new blank page
taco_word.add_paragraph('Random Taco Cookbook', 'Title')
# adding tittle for the recipe

taco_word.add_picture('modified_taco.jpg')
# adding image on word

taco_word.add_paragraph("Tai's Captures")
# going to print this under the picture

taco_word.add_paragraph('https://taco-1150.herokuapp.com/random/?full_taco=true')
# this the url where i got the picture from

taco_word.add_paragraph('Mohamed Hussein') # created by me

taco_word.add_page_break() # this is going to help me break the pages
# for all five components of the recipe

# my next step having all five components of the recipe. and loop over them
for item in range(len(taco_book)): # this is going to help me pull the five recipes
    taco_word.add_paragraph(taco_book[item]["seasoning"]["name"], 'heading 1')
    # first three helping me print the seasoning recipes

    taco_word.add_paragraph(taco_book[item]["seasoning"]["recipe"])
    # this one is for the seasoning recipes

    taco_word.add_paragraph(taco_book[item]["condiment"]["recipe"])
    # item is going to be the condiment recipes

    taco_word.add_paragraph(taco_book[item]["mixin"]["recipe"])
    # this item is going to be mixin recipes

    taco_word.add_paragraph(taco_book[item]["base_layer"]["recipe"])
    # this item is going to be base layer recipes

    taco_word.add_page_break()

taco_word.save('Lo.docx') # saving the word document